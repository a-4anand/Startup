from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
import os, json, re, base64
from dotenv import load_dotenv
import google.generativeai as genai
from PyPDF2 import PdfReader
from django.conf import settings
from pdfminer.high_level import extract_text
from io import BytesIO
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.core.mail import send_mail
import random
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
import google.api_core.exceptions
import razorpay
from django.views.decorators.csrf import csrf_exempt
from .models import Subscription, Contact
from django import forms
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
from django.db.models import F
from django.urls import reverse
from docx import Document
from docx.shared import Inches
from reportlab.lib.styles import ParagraphStyle

load_dotenv()
genai.configure(api_key=settings.GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-pro', generation_config={"temperature": 0.3})
flash_model = genai.GenerativeModel('gemini-1.5-flash', generation_config={"temperature": 0.3})
import re

# --- ADD THIS HELPER FUNCTION at the top of your views.py, after the imports ---
def strip_html_tags(text):
    """A simple function to remove all HTML/XML tags from a string."""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)


def send_otp_email(email, otp):
    subject = "Your OTP for Registration"
    message = f"Your OTP for registration is {otp}. Please do not share this with anyone."
    from_email = "ad3810242@example.com"  # Update with your email
    recipient_list = [email]
    send_mail(subject, message, from_email, recipient_list)
def user_register(request):
    if request.user.is_authenticated:
        messages.info(request, "You are already Registered")
        return redirect('index')

    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            # Don't save yet! Just keep data in session
            email = form.cleaned_data['email']
            username = form.cleaned_data['username']
            password1 = form.cleaned_data['password1']
            password2 = form.cleaned_data['password2']

            # Generate OTP
            otp = random.randint(100000, 999999)
            send_otp_email(email, otp)

            # Save in session
            request.session['otp'] = otp
            request.session['email'] = email
            request.session['username'] = username
            request.session['password1'] = password1
            request.session['password2'] = password2

            return redirect('otp_verify')
        else:
            return render(request, 'Startup_app/register.html', {"form": form, "error": form.errors})

    else:
        form = CustomUserCreationForm()

    return render(request, 'Startup_app/register.html', {"form": form})

def otp_verify_view(request):
    if request.method == 'POST':
        otp_entered = request.POST.get('otp')
        stored_otp = request.session.get('otp')

        if not otp_entered:
            return render(request, 'Startup_app/password/otp-verify.html', {"error": "OTP is required"})
        if otp_entered != str(stored_otp):
            return render(request, 'Startup_app/password/otp-verify.html', {"error": "Invalid OTP. Please try again."})

        # Retrieve data from session
        email = request.session.get('email')
        username = request.session.get('username')
        password1 = request.session.get('password1')
        password2 = request.session.get('password2')

        # Check if email already exists
        if User.objects.filter(email=email).exists():
            return render(request, 'Startup_app/password/otp-verify.html', {"error": "This email is already registered."})

        # Create user now
        user = User.objects.create_user(username=username, email=email, password=password1)

        login(request, user)

        # Clear session
        for key in ['otp', 'email', 'username', 'password1', 'password2']:
            request.session.pop(key, None)

        return redirect('index')

    return render(request, 'Startup_app/password/otp-verify.html')

def user_login(request):
    if request.user.is_authenticated:
        messages.info(request, "You are already logged in")
        return redirect('index')

    if request.method == "POST":
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, "Login successful!")
            return redirect('index')
        else:
            messages.error(request, "Invalid username or password!")
    return render(request, "Startup_app/login.html")
def user_logout(request):
    logout(request)
    messages.success(request, "Logged out successfully!")
    return redirect("index")
@login_required
def profile_view(request):
    subscription = None
    generations_left = 0
    try:
        subscription = Subscription.objects.get(user=request.user)
        if subscription.is_paid:
            # Calculate remaining generations
            generations_left = 5 - subscription.download_count
            if generations_left < 0:
                generations_left = 0
    except Subscription.DoesNotExist:
        subscription = None

    return render(request, "Startup_app/profile.html", {
        "subscription": subscription,
        "generations_left": generations_left
    })
def index(request):
    return render(request, 'Startup_app/index.html')
def form_view(request):
    # Clear previous results from session when visiting the form page
    request.session.pop('ats_score', None)
    request.session.pop('comments', None)
    request.session.pop('generated_resume', None)
    return render(request, 'Startup_app/form.html')
def extract_text_from_pdf(uploaded_file):
    try:
        # Use BytesIO to handle the in-memory file
        file_bytes = BytesIO(uploaded_file.read())
        text = extract_text(file_bytes)
        # Reset pointer for any subsequent reads if necessary, though not needed here
        uploaded_file.seek(0)
        return text
    except Exception as e:
        print(f"Error extracting text: {e}") # For debugging
        return None
def analyze_resume(request):
        if request.method == "POST":
            resume_file = request.FILES.get("resume")
            job_description = request.POST.get("job_description")
            experience_level = request.POST.get("experience_level")
            if not resume_file or not job_description or not experience_level:
                messages.error(request,
                               "Please fill out all fields: upload resume, select experience, and paste job description.")
                return redirect("form_view")

            resume_text = extract_text_from_pdf(resume_file)
            if not resume_text:
                messages.error(request, "Failed to read text from the uploaded PDF. Please try another file.")
                return redirect("form_view")

            # --- Store the original resume text in the session for later use ---
            request.session['original_resume_text'] = resume_text
            request.session['job_description'] = job_description
            request.session['experience_level'] = experience_level

            # --- Get ONLY the Initial ATS Score (The Free Part) ---
            initial_scoring_prompt = f"""
You are an elite, enterprise-grade ATS (Applicant Tracking System) Scoring Engine. Your function is to perform a rigorous, quantitative analysis of a candidate's resume against a target job description. You must operate with extreme precision and adhere strictly to the multi-step protocol outlined below.

---
**SCORING RUBRIC (100 points total):**
You will calculate the final score based on the following weighted categories:

1.  **Keyword & Skill Alignment (40 points):** Direct match of essential skills, technologies, and keywords mentioned in the job description.
2.  **Relevant Experience (30 points):** Alignment of job titles, years of experience, and industry context with the role's requirements.
3.  **Quantifiable Achievements (20 points):** Presence and quality of metrics, numbers, and data-driven results in the experience section.
4.  **Education & Certifications (10 points):** Match of required degrees, certifications, and educational background.

---
**MULTI-STEP ANALYSIS PROTOCOL:**
Before providing any output, you must perform the following analysis internally:

* **Step 1: Deconstruction:** Identify the top 10 most critical hard skills, tools, and qualifications from the **Job Description**.
* **Step 2: Mapping:** Scan the **Candidate's Resume** and list which of the critical factors from Step 1 are present.
* **Step 3: Scoring:** Based on the mapping in Step 2 and a holistic review, assign a score for each of the 4 categories in the **Scoring Rubric**.
* **Step 4: Synthesis:** Sum the scores from Step 3 to calculate the final, total ATS score.
* **Step 5: Feedback Generation:** Use the results of your analysis to generate feedback. The two strengths MUST come from the highest-scoring categories. The two weaknesses MUST come from the lowest-scoring categories. The feedback must be specific and actionable.

---
**INPUTS:**

**1. Job Description:** 
{job_description}
**2. Candidate's Resume:**

{resume_text}


---
**STRICT OUTPUT FORMAT:**
Your entire output must be ONLY the following. Do not add any other text, explanations, or conversational filler.

* **Line 1:** The final ATS score, formatted exactly as: `SCORE: [Total Score]`
* **Line 2:** The first strength bullet point, starting with `* Strength:`.
* **Line 3:** The second strength bullet point, starting with `* Strength:`.
* **Line 4:** The first weakness bullet point, starting with `* Weakness:`.
* **Line 5:** The second weakness bullet point, starting with `* Weakness:`.

"""
            try:
                initial_response = flash_model.generate_content(initial_scoring_prompt)
                raw_initial_output = initial_response.text.strip()

                # More robust parsing for the score
                lines = raw_initial_output.splitlines()
                initial_first_line = lines[0].strip() if lines else ""

                # This regex now looks for any number in the first line
                initial_match = re.search(r'\d+', initial_first_line)

                initial_ats_score = initial_match.group(0) if initial_match else "N/A"
                initial_comments = "\n".join(lines[1:]).strip()

            except Exception as e:
                messages.error(request, f"An error occurred during analysis: {e}")
                print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
                print(f"CRITICAL ERROR IN GEMINI API CALL: {e}")
                print(f"ERROR TYPE: {type(e)}")
                print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
                return redirect('form_view')

            # --- Render the initial results page ---
            return render(request, "Startup_app/form.html", {
                "ats_score": initial_ats_score,
                "comments": initial_comments,
            })

        # If GET request, redirect to the form
        return redirect("form_view")


# --- THIS IS THE FINAL, CORRECTED VERSION. PLEASE REPLACE YOUR ENTIRE FUNCTION WITH THIS. ---
@login_required
def generate_final_resume(request):

    try:
        subscription = Subscription.objects.get(user=request.user)
        if not subscription.is_paid:
            messages.error(request, "You must complete the payment to generate the final resume.")
            return redirect("upgrade_page")

        if subscription.download_count >= 5:
            messages.error(request,
                           "You have reached your limit of 5 resume generations. Please purchase again to continue.")

            # IMPORTANT: Expire their current plan to force re-payment
            subscription.is_paid = False
            subscription.save()

            return redirect("upgrade_page")
    except Subscription.DoesNotExist:
        messages.error(request, "No subscription found. Please complete the payment process.")
        return redirect("upgrade_page")

    resume_text = request.session.get('original_resume_text')
    job_description = request.session.get('job_description')
    experience_level = request.session.get('experience_level')
    if not all([resume_text, job_description, experience_level]):
        messages.error(request,
                       "Your session has expired or is missing key information. Please analyze your resume again.")
        return redirect("form_view")

    # --- Step 1: Rewrite the Resume (Same prompt as before) ---

    # In generate_final_resume, replace your entire rewrite_prompt with this:
    rewrite_prompt = f"""
    You are "Zeus," an elite AI Career Strategist and Document Architect. Your function transcends simple writing; you are to build a compelling career narrative that functions as a high-conversion marketing document for the candidate.

    **PRIME DIRECTIVE:**
    Deconstruct the provided inputs and synthesize a "Tier-1" resume. This document must be engineered to bypass 99.9% of all Applicant Tracking Systems (ATS) and, more importantly, to immediately capture the attention and confidence of a human hiring manager.

    **CORE METHODOLOGY: The Impact Delta**
    Your primary goal is not to list tasks, but to demonstrate the "Impact Delta"—the measurable change and value the candidate brought to each role. Every bullet point must answer the silent question of the recruiter: "So what?"

    **STRATEGIC EXECUTION FRAMEWORK:**

    1.  **Cultural Resonance Analysis:** Before writing, analyze the job description for cultural keywords (e.g., "fast-paced," "collaborative," "data-driven," "innovative"). The TONE of the resume must subtly mirror this culture. A resume for a creative agency should feel different from one for a bank.

    2.  **The 6-Second Test Optimization:** Structure the top third of the resume for maximum information density. A recruiter must grasp the candidate's core value proposition within a 6-second glance. This is achieved through a powerful header, a concise Value Proposition Statement, and a well-organized Skills Matrix.

    3.  **Impact-Driven Narrative Flow:** The experience section must not be a list of jobs. It must tell a story of growth. Structure bullet points to show a clear progression of skills, responsibilities, and quantifiable impact over time.

    ---
    **SECTION-SPECIFIC MANDATES:**

    * **CANDIDATE HEADER:**
        * Extract and prominently display the candidate's Full Name, Phone, Email, and a clickable LinkedIn URL. The formatting must be clean, modern, and professional.

    * **VALUE PROPOSITION STATEMENT:**
        * Compose a 3-4 line strategic narrative. It must begin with a powerful, role-specific title (e.g., "Senior Data Scientist with 8+ years of experience in predictive modeling"). It must seamlessly integrate the top skills from the job description and showcase the candidate's most impressive, quantifiable career achievement.

    * **HYBRID SKILLS MATRIX:**
        * This is not a keyword dump. Create logical, clean categories (`Technical Proficiencies:`, `Software & Tools:`, `Certifications:`). Prioritize the skills listed in the job description.

    * **PROFESSIONAL EXPERIENCE:**
        * Frame each job as a mission.
        * Each bullet point MUST be a high-impact, quantified achievement that starts with a powerful action verb (e.g., Architected, Spearheaded, Accelerated, Revitalized).
        * Every metric and number MUST be bolded using markdown (`**+25%**`, `**$500K**`).

    ---
    **NON-NEGOTIABLE PROTOCOLS:**

    * **Output Format:** ONLY the rewritten resume text.
    * **Visual Separator:** Use `===== HORIZONTAL_RULE =====`.
    * **Section Order:** CANDIDATE HEADER, VALUE PROPOSITION STATEMENT, HYBRID SKILLS MATRIX, PROFESSIONAL EXPERIENCE, PROJECTS, EDUCATION.
    * **Forbidden Elements:** NO personal pronouns, NO clichés, NO generic soft skills (e.g., "team player" unless it's a key term in the JD). DO NOT use the words 'Challenge', 'Action', or 'Result'.

    ---
    **INPUTS:**
    1.  **CANDIDATE PROFILE:** {experience_level}
    2.  **TARGET JOB DESCRIPTION:** --- {job_description} ---
    3.  **ORIGINAL RESUME CONTENT:** --- {resume_text} ---

    Execute the Prime Directive.
    """

    rewrite_response = model.generate_content(rewrite_prompt)
    generated_resume = rewrite_response.text.strip()
    request.session["generated_resume"] = generated_resume



    subscription.download_count = F('download_count') + 1
    subscription.save()


    # --- Step 3: Create PDF for Preview (Same PDF logic as before) ---
    pdf_buffer = BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=LETTER, leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72)
    primary_color = colors.HexColor("#0D47A1")  # A professional deep blue
    body_text_color = colors.HexColor("#2C2C2C")

    heading_style = ParagraphStyle(
            name='Heading2',
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=14,  # Line spacing
            textColor=primary_color,
            spaceAfter=6
        )

    body_style = ParagraphStyle(
            name='Normal',
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            textColor=body_text_color
        )


    story = []

    for line in generated_resume.split('\n'):
        # Convert markdown bold (**text**) to reportlab bold (<b>text</b>)
        line_with_formatting = re.sub(r'\*\*(.*?)\*\*', '<b>\\1</b>', line)

        if not line_with_formatting.strip():
            continue

        if "===== HORIZONTAL_RULE =====" in line_with_formatting:
            story.append(Spacer(1, 8))
            story.append(Table([[""]], colWidths=[doc.width],
                               style=TableStyle([('LINEBELOW', (0, 0), (-1, -1), 1, primary_color)])))
            story.append(Spacer(1, 8))
        elif line_with_formatting.strip().isupper() and len(line_with_formatting.strip().split()) < 5:
            story.append(Paragraph(line_with_formatting.strip(), heading_style))
            story.append(Spacer(1, 2))  # Reduced spacer for tighter look
        elif line_with_formatting.strip().startswith(("•", "-", "*")):
            story.append(Paragraph("• " + line_with_formatting.lstrip(" -*").strip(), body_style))
        else:
            story.append(Paragraph(line_with_formatting, body_style))

    doc.build(story)
    pdf_data = pdf_buffer.getvalue()
    pdf_buffer.close()
    generated_resume_pdf = base64.b64encode(pdf_data).decode("utf-8")

    # --- Render the FINAL results page with PDF preview and download ---
    return render(request, "Startup_app/form.html", {

        "generated_resume_pdf": generated_resume_pdf
    })
@login_required
def download_resume_pdf(request):
    # The generation was the limited action. Download is now free for generated content.
    resume_content = request.session.get("generated_resume")
    if not resume_content:
        messages.error(request, "Could not find the generated resume in your session. Please generate it again.")
        return redirect("form_view")

    # --- PDF Generation Logic (your existing code is fine) ---
    pdf_buffer = BytesIO()
    # ... (your code to build the PDF document) ...
    doc = SimpleDocTemplate(pdf_buffer, pagesize=LETTER, leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []
    for line in resume_content.split("\n"):
        clean_line = re.sub(r"[*_]{1,2}(.*?)\1?[*_]{1,2}", r"\1", line).strip()
        if "===== HORIZONTAL_RULE =====" in clean_line:
            story.append(Spacer(1, 8))
            story.append(Table([[""]], colWidths=[450],
                               style=TableStyle([('LINEBELOW', (0, 0), (-1, -1), 1, colors.HexColor("#0D47A1"))])))
            story.append(Spacer(1, 8))
        elif clean_line.isupper() and len(clean_line.split()) < 5:
            story.append(Paragraph(clean_line, styles["h2"]))
            story.append(Spacer(1, 4))
        elif clean_line.startswith(("-", "*")):
            story.append(Paragraph("• " + clean_line.lstrip("-* ").strip(), styles["Normal"]))
        elif clean_line:
            story.append(Paragraph(clean_line, styles["Normal"]))
    doc.build(story)
    pdf_data = pdf_buffer.getvalue()
    pdf_buffer.close()

    response = HttpResponse(pdf_data, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="ATS_Optimized_Resume.pdf"'
    return response
@login_required
def download_resume_word(request):
    # 1. Get the resume content from the session
    resume_content = request.session.get("generated_resume")
    if not resume_content:
        messages.error(request, "Could not find the generated resume in your session. Please generate it again.")
        return redirect("form_view")

    # 2. Create a new Word document
    document = Document()

    # Optional: Set margins for a professional look
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # 3. Parse the resume text and add it to the document
    for line in resume_content.split("\n"):
        clean_line = re.sub(r"[*_]{1,2}(.*?)\1?[*_]{1,2}", r"\1", line).strip()

        if "===== BLUE LINE =====" in clean_line:
            # Add a simple paragraph for spacing, as a visual line is complex in .docx
            document.add_paragraph()
        elif clean_line.isupper() and len(clean_line.split()) < 5 and clean_line:
            # Add section headings like "EXPERIENCE"
            document.add_heading(clean_line, level=2)
        elif clean_line.startswith(("-", "*")):
            # Add bullet points
            document.add_paragraph(clean_line.lstrip("-* ").strip(), style='List Bullet')
        elif clean_line:
            # Add regular text
            document.add_paragraph(clean_line)

    # 4. Save the document to an in-memory buffer
    doc_buffer = BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)

    # 5. Create the HTTP response to send the file
    response = HttpResponse(
        doc_buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="ATS_Optimized_Resume.docx"'
    return response


from django.contrib import messages
from .models import Contact

def contact_view(request):
    if request.method == "POST":
        first_name = request.POST.get("first_name")
        last_name = request.POST.get("last_name")
        email = request.POST.get("email")
        phone = request.POST.get("phone")
        subject = request.POST.get("subject")
        message = request.POST.get("message")

        # Save to DB
        Contact.objects.create(
            first_name=first_name,
            last_name=last_name,
            email=email,
            phone=phone,
            subject=subject,
            message=message
        )

        messages.success(request, "Thank you! Your message has been sent successfully.")
        return redirect("contact")

    return render(request, "Startup_app/contact.html")




import razorpay
from django.views.decorators.csrf import csrf_exempt

# In your views.py file

@login_required(login_url='/login/')
def upgrade_page(request):
    # Get or create a subscription object for the logged-in user
    subscription, created = Subscription.objects.get_or_create(user=request.user)

    # --- KEY CHANGE: Check if the user is already subscribed ---
    if subscription.is_paid:
        messages.info(request, "You already have an active premium subscription! 🚀")
        return redirect('profile') # Redirect them away from the payment page

    # --- If not subscribed, proceed with creating a payment order ---
    try:
        client = razorpay.Client(auth=(settings.RAZORPAY_KEY_ID, settings.RAZORPAY_KEY_SECRET))

        # Note: I've corrected the amount to 4900 paise (₹49) to match your template
        order = client.order.create({
            "amount": 4900,
            "currency": "INR",
            "payment_capture": 1
        })

        # Store the new intended order ID
        subscription.intended_razorpay_order_id = order["id"]
        subscription.save()

        context = {
            "order_id": order["id"],
            "razorpay_key": settings.RAZORPAY_KEY_ID,
            "subscription": subscription # Pass subscription status to the template
        }
        return render(request, "Startup_app/plan.html", context)

    except Exception as e:
        print(f"Error creating Razorpay order: {e}")
        messages.error(request, "Could not connect to the payment gateway. Please try again later.")
        return redirect('index')
@csrf_exempt

def payment_success(request):
    if request.method == "POST":
        data = json.loads(request.body)
        payment_id = data.get("razorpay_payment_id")
        order_id = data.get("razorpay_order_id")
        signature = data.get("razorpay_signature")

        client = razorpay.Client(auth=(settings.RAZORPAY_KEY_ID, settings.RAZORPAY_KEY_SECRET))
        try:
            client.utility.verify_payment_signature(data)
        except:
            return JsonResponse({"status": "failed", "message": "Signature verification failed."})

        try:
            if not request.user.is_authenticated:
                return JsonResponse({"status": "failed", "message": "User not authenticated."})
            subscription = Subscription.objects.get(user=request.user, intended_razorpay_order_id=order_id)
            subscription.razorpay_order_id = order_id
            subscription.razorpay_payment_id = payment_id
            subscription.razorpay_signature = signature
            subscription.plan = "Purchase"
            subscription.is_paid = True
            subscription.download_count = 0
            subscription.save()

            # *** KEY CHANGE: Return a URL for redirection ***
            return JsonResponse({
                "status": "success",
                "redirect_url": "/generate-final-resume/" # The URL to our new view
            })
        except Subscription.DoesNotExist:
            return JsonResponse({"status": "failed", "message": "Subscription not found."})

class CustomUserCreationForm(UserCreationForm):
    email = forms.EmailField(required=True)

    class Meta:
        model = User
        fields = ("username", "email", "password1", "password2")

    def save(self, commit=True):
        user = super().save(commit=False)
        user.email = self.cleaned_data["email"]
        if commit:
            user.save()
        return user
